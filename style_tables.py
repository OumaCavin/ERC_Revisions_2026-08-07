#!/usr/bin/env python3
"""
Apply professional table styling to a docx document.
Adds light blue borders, bold headers, and proper formatting.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
import re

def hex_to_rgb(hex_color):
    """Convert hex color to RGB tuple."""
    hex_color = hex_color.lstrip('#')
    return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))

def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    """
    Set individual cell borders.
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # Remove existing borders
    for border_name in ['w:top', 'w:left', 'w:bottom', 'w:right']:
        for border in tcPr.findall(qn(border_name)):
            tcPr.remove(border)

    # Set new borders
    borders_map = {
        'top': top,
        'bottom': bottom,
        'left': left,
        'right': right
    }

    for edge_name, border in borders_map.items():
        if border:
            tag = f'w:{edge_name}'
            element = OxmlElement(tag)
            element.set(qn('w:val'), border.get('val', 'single'))
            element.set(qn('w:sz'), str(border.get('sz', 4)))
            element.set(qn('w:space'), '0')
            element.set(qn('w:color'), border.get('color', '000000'))
            tcPr.append(element)

def set_cell_shading(cell, fill_color):
    """
    Set cell background shading.
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # Remove existing shading
    for shading in tcPr.findall(qn('w:shd')):
        tcPr.remove(shading)

    # Add new shading
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), fill_color)
    tcPr.append(shading)

def set_cell_margins(cell, margin=50):
    """
    Set cell margins.
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # Remove existing margins
    for margins in tcPr.findall(qn('w:tcMar')):
        tcPr.remove(margins)

    # Add new margins
    margins = OxmlElement('w:tcMar')
    for edge in ['top', 'left', 'bottom', 'right']:
        element = OxmlElement(f'w:{edge}')
        element.set(qn('w:w'), str(margin))
        element.set(qn('w:type'), 'dxa')
        margins.append(element)
    tcPr.append(margins)

def style_table(table, is_header_row=True):
    """
    Apply professional table styling.
    """
    # Colors
    LIGHT_BLUE = 'C5D7E8'
    HEADER_BLUE = '1F5C8B'
    BORDER_BLUE = '4A90C2'

    # Style each cell
    for row_idx, row in enumerate(table.rows):
        is_header = (is_header_row and row_idx == 0)

        for cell in row.cells:
            # Set cell margins
            set_cell_margins(cell, margin=80)

            # Set borders
            border_color = HEADER_BLUE if is_header else BORDER_BLUE
            border_width = 8 if is_header else 4

            set_cell_border(cell,
                top={'val': 'single', 'sz': border_width, 'color': border_color},
                bottom={'val': 'single', 'sz': border_width, 'color': border_color},
                left={'val': 'single', 'sz': border_width, 'color': border_color},
                right={'val': 'single', 'sz': border_width, 'color': border_color}
            )

            # Set shading
            if is_header:
                set_cell_shading(cell, LIGHT_BLUE)
            else:
                set_cell_shading(cell, 'FFFFFF')

            # Set text formatting
            for paragraph in cell.paragraphs:
                # Set alignment
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

                for run in paragraph.runs:
                    # Set font
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(10)

                    if is_header:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(*hex_to_rgb(HEADER_BLUE))
                    else:
                        run.font.bold = False
                        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    # Set table properties
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    # Set table width
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = OxmlElement('w:tblW')
        tblPr.append(tblW)
    tblW.set(qn('w:w'), '5000')
    tblW.set(qn('w:type'), 'pct')

    # Set table borders (outer border)
    tblBorders = tblPr.find(qn('w:tblBorders'))
    if tblBorders is None:
        tblBorders = OxmlElement('w:tblBorders')
        tblPr.append(tblBorders)

    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = tblBorders.find(qn(f'w:{border_name}'))
        if border is None:
            border = OxmlElement(f'w:{border_name}')
            tblBorders.append(border)
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '8')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), BORDER_BLUE)

def process_document(input_file, output_file):
    """
    Process the document and apply table styling.
    """
    print(f"Processing: {input_file}")

    doc = Document(input_file)

    table_count = 0
    for table in doc.tables:
        table_count += 1
        style_table(table)
        print(f"  Styled table {table_count}")

    doc.save(output_file)
    print(f"Saved: {output_file}")
    return table_count

if __name__ == '__main__':
    input_file = '/workspace/ERC_Revisions_2026-08-07/Research_proposal_styled.docx'
    output_file = '/workspace/ERC_Revisions_2026-08-07/Research_proposal_styled.docx'

    count = process_document(input_file, output_file)
    print(f"\nStyled {count} tables successfully!")
