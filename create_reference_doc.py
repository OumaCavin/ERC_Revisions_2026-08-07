#!/usr/bin/env python3
"""
Create a pandoc reference document with styled tables.
This creates a custom reference.docx with light blue borders and bold headers.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_border(cell, **kwargs):
    """
    Set cell borders with specified color and width.
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    for edge in ['top', 'left', 'bottom', 'right']:
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # Create the border element
            element = OxmlElement(tag)
            element.set(qn('w:val'), edge_data.get('val', 'single'))
            element.set(qn('w:sz'), str(edge_data.get('sz', 4)))  # Border width (in 1/8 pt)
            element.set(qn('w:space'), '0')
            element.set(qn('w:color'), edge_data.get('color', '000000'))

            tcPr.append(element)

def set_cell_shading(cell, fill_color):
    """
    Set cell background shading.
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), fill_color)
    tcPr.append(shading)

def style_table(table):
    """
    Apply professional table styling with light blue borders.
    """
    LIGHT_BLUE = 'C5D7E8'  # Light blue color
    DARK_BLUE = '1F5C8B'   # Darker blue for header text

    # Set table borders
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell,
                top={'val': 'single', 'sz': 8, 'color': LIGHT_BLUE},
                bottom={'val': 'single', 'sz': 8, 'color': LIGHT_BLUE},
                left={'val': 'single', 'sz': 8, 'color': LIGHT_BLUE},
                right={'val': 'single', 'sz': 8, 'color': LIGHT_BLUE}
            )

    # Style header row
    if table.rows:
        header_row = table.rows[0]
        for cell in header_row.cells:
            set_cell_shading(cell, LIGHT_BLUE)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0x1F, 0x5C, 0x8B)  # Dark blue
                    run.font.size = Pt(10)
                    run.font.name = 'Times New Roman'
        # Add bottom border to header
        for cell in header_row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            bottom_border = OxmlElement('w:bottom')
            bottom_border.set(qn('w:val'), 'single')
            bottom_border.set(qn('w:sz'), '12')  # Thicker border
            bottom_border.set(qn('w:space'), '0')
            bottom_border.set(qn('w:color'), DARK_BLUE)
            tcPr.append(bottom_border)

    # Style data rows
    for row in table.rows[1:]:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    run.font.name = 'Times New Roman'
                    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

# Create a new document
doc = Document()

# Add a title
title = doc.add_heading('Reference Document for Pandoc Conversion', 0)
title.alignment = WD_TABLE_ALIGNMENT.CENTER

# Add description
doc.add_paragraph('This document serves as a reference for pandoc conversions with styled tables.')

# Create sample table 1: Simple two-column table
doc.add_heading('Sample Table 1', level=1)
table1 = doc.add_table(rows=4, cols=2)
table1.style = 'Table Grid'
table1.alignment = WD_TABLE_ALIGNMENT.CENTER

# Header row
hdr_cells = table1.rows[0].cells
hdr_cells[0].text = 'Column 1'
hdr_cells[1].text = 'Column 2'

# Data rows
data = [
    ('Item 1', 'Value A'),
    ('Item 2', 'Value B'),
    ('Item 3', 'Value C'),
]
for i, (col1, col2) in enumerate(data):
    row_cells = table1.rows[i+1].cells
    row_cells[0].text = col1
    row_cells[1].text = col2

style_table(table1)

# Create sample table 2: Three-column table
doc.add_heading('Sample Table 2', level=1)
table2 = doc.add_table(rows=5, cols=3)
table2.style = 'Table Grid'
table2.alignment = WD_TABLE_ALIGNMENT.CENTER

hdr_cells = table2.rows[0].cells
hdr_cells[0].text = 'ID'
hdr_cells[1].text = 'Description'
hdr_cells[2].text = 'Status'

data2 = [
    ('1', 'First item', 'Active'),
    ('2', 'Second item', 'Pending'),
    ('3', 'Third item', 'Completed'),
    ('4', 'Fourth item', 'Inactive'),
]
for i, (id_val, desc, status) in enumerate(data2):
    row_cells = table2.rows[i+1].cells
    row_cells[0].text = id_val
    row_cells[1].text = desc
    row_cells[2].text = status

style_table(table2)

# Add paragraph spacing
doc.add_paragraph()

# Save the document as reference.docx
doc.save('/workspace/ERC_Revisions_2026-08-07/reference.docx')
print("Reference document created successfully: reference.docx")
